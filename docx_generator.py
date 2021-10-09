# -*- coding: utf-8 -*-

'''
Implement floating image based on python-docx.

- Text wrapping style: BEHIND TEXT <wp:anchor behindDoc="1">
- Picture position: top-left corner of PAGE `<wp:positionH relativeFrom="page">`.

Create a docx sample (Layout | Positions | More Layout Options) and explore the 
source xml (Open as a zip | word | document.xml) to implement other text wrapping
styles and position modes per `CT_Anchor._anchor_xml()`.
'''

from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.shared import Cm

import datetime

# refer to docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )


# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename    
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

# refer to docx.oxml.shape.__init__.py
register_element_cls('wp:anchor', CT_Anchor)


if __name__ == '__main__':

    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()

    array_index = [] # номерки всех мишеней в документе

# следующий блок кода нужно будет перевести на GUI
#######################################################################
    weapon = input("введите какое оружие - винтовка/пистолет: ")
    exercise = input("введите какое упражнение - 20/30/40/60: ")
    if(weapon == "винтовка" and exercise == "40") or (weapon == "винтовка" and exercise == "60"):
        exercise_mode = int(input("введите по сколько выстрелов будет - 1/2: "))
    else:
        exercise_mode = 5
    shifts = int(input("введите сколько будет смен: "))
    #shields = int(input("введите сколько будет щитов: "))
    shields = [int(s) for s in input("введите номера щитов через запятую: ").split(",")]
    number_of_targets = int(int(exercise)/int(exercise_mode))  # количество мишеней для одного человека - последняя цифра в коде

    print("--------------------------------")
    print("собранная информация:")
    print("оружие: " + weapon)
    print("упражнение: " + exercise)
    print("выстрелов: " + str(exercise_mode))
    print("смен: " + str(shifts))
    print("щиты: " + str(shields))
    print("количество мишеней для одного человека: " + str(number_of_targets))
    print("--------------------------------")

    for num_shifts in range(shifts + 1):  # проходимся по каждой смене
        if num_shifts == 0:  # пропускаем нулевой элемент, чтобы везде начинать с 1
            continue
        for num_shields in shields:  # проходимся по каждому щиту
            if num_shields == 0:
                continue
            for num_targets in range(number_of_targets + 1):  # проходимся по всем мишеням
                if num_targets == 0:
                    continue
                # формирую индексы
                # первые две цифры - номер смены
                if len(str(num_shifts)) < 2:
                    first_index = "0" + str(num_shifts)
                else:
                    first_index = str(num_shifts)
                # вторые две цифры - номер щита
                if len(str(num_shields)) < 2:
                    second_index = "0" + str(num_shields)
                else:
                    second_index = str(num_shields)
                # третьи две цифры - номер мишени
                if len(str(num_targets)) < 2:
                    third_index = "0" + str(num_targets)
                else:
                    third_index = str(num_targets)
                
                index = first_index + second_index + third_index
                print(index)
                array_index.append(index)

#######################################################################

    if(weapon == "винтовка"):
        offsetX = 70
        offsetY = 55
        p = document.add_paragraph()
        # paragraph_format = p.paragraph_format
        # paragraph_format.space_after = Pt(0)
        # paragraph_format.line_spacing = 1.75
        
        chetchik = 0
        chetchik_page = -1

        for i in array_index:
            
            chetchik += 1

            if(chetchik <= 2):
                p.add_run('                            ')
                run = p.add_run(i)  # тут пишется наш сгенерированый номер
                run.font.name = 'Calibri'
                run.font.size = Pt(18)

                if(chetchik < 2):
                    p.add_run('                                          ')
            else:
                chetchik = 1
                p = document.add_paragraph("")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 14.9
                p = document.add_paragraph()
                p.add_run('                           ')
                run = p.add_run(i)  # тут пишется наш сгенерированый номер
                run.font.name = 'Calibri'
                run.font.size = Pt(18)
                p.add_run('                                            ')

            # следущий блок кода вставляет 6 картинок маленькой мишени каждый шестой номер (чтобы они были на каждой странице)
            ####################################
            chetchik_page += 1
            if(chetchik_page == 6) or (chetchik_page == 0):
                chetchik_page = 0
                add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(1 + offsetX), pos_y=Pt(1 + offsetY))  # x - ширина, y - высота
                add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(1 + offsetX), pos_y=Pt(227.5 + offsetY))
                add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(1 + offsetX), pos_y=Pt(227 * 2 + offsetY))
                add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(227.5 + offsetX), pos_y=Pt(1 + offsetY))
                add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(227.5 + offsetX), pos_y=Pt(227.5 + offsetY))
                add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(227.5 + offsetX), pos_y=Pt(227 * 2 + offsetY))
            ####################################
    
    if(weapon == "пистолет"):
        offsetX = 70
        offsetY = 68
        p = document.add_paragraph()
        
        chetchik = 0

        for i in array_index:
            # следущий блок кода вставляет 1 картинку большой мишени и 2 маленькой каждый номер (чтобы они были на каждой странице)
            ####################################
            add_float_picture(p, 'мишень скрин 17х17см.png', width=Cm(17.0), pos_x=Pt(1 + offsetX), pos_y=Pt(1 + 40))  # x - ширина, y - высота
            add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(1 + offsetX), pos_y=Pt(227 * 2 + offsetY))
            add_float_picture(p, 'мишень скрин 8х8см.png', width=Cm(8.0), pos_x=Pt(227.5 + offsetX), pos_y=Pt(227 * 2 + offsetY))
            ####################################
            chetchik += 1

            p.add_run('')
            run = p.add_run(i)  # тут пишется наш сгенерированый номер
            run.font.name = 'Calibri'
            run.font.size = Pt(28)
            
            if chetchik < len(array_index): # чтобы после последнего номера не было отступа
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 30
                p = document.add_paragraph()

    if(weapon == "винтовка"):
        weapon_code = "ВП"
    if(weapon == "пистолет"):
        weapon_code = "ПП"

    now = datetime.datetime.now()
    time = now.strftime("%d.%m.%Y_%H.%M")
    document.save(weapon_code + str(exercise) + " " + "смен " + str(shifts) + " " + "щиты " + str(shields) + " " + time + ".docx")
